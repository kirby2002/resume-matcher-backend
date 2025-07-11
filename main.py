# from fastapi import FastAPI, File, UploadFile, Form
# from fastapi.middleware.cors import CORSMiddleware
# import pdfplumber
# import docx
# from sentence_transformers import SentenceTransformer, util

# app = FastAPI()

# # Allow frontend access
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # Allow all origins
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Load AI model
# model = SentenceTransformer('./all-MiniLM-L6-v2')

# # Extract text from uploaded resume
# def extract_text(file: UploadFile):
#     if file.filename.endswith(".pdf"):
#         with pdfplumber.open(file.file) as pdf:
#             return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#     elif file.filename.endswith(".docx"):
#         doc = docx.Document(file.file)
#         return "\n".join([para.text for para in doc.paragraphs])
#     else:
#         return ""

# @app.post("/analyze")
# async def analyze(resume: UploadFile = File(...), jd: str = Form(...)):
#     resume_text = extract_text(resume)
#     embeddings = model.encode([resume_text, jd])
#     similarity = util.cos_sim(embeddings[0], embeddings[1]).item()
#     return {
#         "similarity": round(similarity * 100, 2),
#         "match": "High" if similarity > 0.7 else "Moderate" if similarity > 0.4 else "Low"
#     }

from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
import pdfplumber
import docx
from io import BytesIO
from sentence_transformers import SentenceTransformer, util

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

model = SentenceTransformer('paraphrase-MiniLM-L3-v2')

# def extract_text(file: UploadFile):
#     contents = file.file.read()
#     if file.filename.endswith(".pdf"):
#         with pdfplumber.open(BytesIO(contents)) as pdf:
#             return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#     elif file.filename.endswith(".docx"):
#         doc = docx.Document(BytesIO(contents))
#         return "\n".join([para.text for para in doc.paragraphs])
#     return ""

def extract_text(file: UploadFile):
    if file.filename.endswith(".pdf"):
        with pdfplumber.open(file.file) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif file.filename.endswith(".docx"):
        doc = docx.Document(file.file)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        return ""
    
    # Remove email, phone, links, and declarations
    cleaned = []
    for line in text.splitlines():
        if any(x in line.lower() for x in ["@", "linkedin", "phone", "declaration", "india"]):
            continue
        if len(line.strip()) > 3:
            cleaned.append(line.strip())
    return "\n".join(cleaned)

@app.post("/analyze")
async def analyze(resume: UploadFile = File(...), jd: str = Form(...)):
    resume_text = extract_text(resume)
    embeddings = model.encode([resume_text, jd])
    similarity = util.cos_sim(embeddings[0], embeddings[1]).item()
    return {
        "similarity": round(similarity * 100, 2),
        "match": "High" if similarity > 0.7 else "Moderate" if similarity > 0.4 else "Low"
    }

import os

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 10000))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)
